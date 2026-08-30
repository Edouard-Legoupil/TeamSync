"""Extract plain text from uploaded transcript files (.txt / .docx / .vtt)."""

from __future__ import annotations

import re
from io import BytesIO

_CUE_TIMING = re.compile(r"-->")
_TAG = re.compile(r"<[^>]+>")


def _vtt_to_text(raw: str) -> str:
    """Strip WebVTT headers, cue timing, and inline tags down to spoken text."""
    lines: list[str] = []
    for line in raw.splitlines():
        line = line.strip()
        if not line:
            continue
        if line.startswith(("WEBVTT", "NOTE", "STYLE", "REGION")):
            continue
        if _CUE_TIMING.search(line):
            continue
        # Drop inline speaker/style tags like `<v Roger>…`.
        line = _TAG.sub("", line).strip()
        if line:
            lines.append(line)
    return "\n".join(lines)


def extract_text(filename: str, content: bytes) -> str:
    """Return the raw text of an uploaded transcript.

    Images and non-text artifacts in ``.docx`` files are ignored: only
    paragraph and table cell text is captured.
    """
    lower = filename.lower()

    if lower.endswith((".txt", ".md")):
        return content.decode("utf-8", errors="replace")

    if lower.endswith(".vtt"):
        return _vtt_to_text(content.decode("utf-8", errors="replace"))

    if lower.endswith(".docx"):
        from docx import Document

        doc = Document(BytesIO(content))
        parts: list[str] = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return "\n".join(parts)

    raise ValueError("Unsupported file type. Please upload a .txt, .vtt, or .docx file.")
