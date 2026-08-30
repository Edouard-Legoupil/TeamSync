"""Convert Markdown meeting content into a professional Word (.docx) document.

Maps Markdown headings to Word Heading styles, bullets/numbered lists to Word
list styles, and Markdown tables to real Word tables.
"""

from __future__ import annotations

import re
from io import BytesIO

from docx import Document
from docx.shared import Inches, Pt, RGBColor

NAVY = RGBColor(0x18, 0x37, 0x5F)
BLUE = RGBColor(0x00, 0x72, 0xBC)
CHARCOAL = RGBColor(0x33, 0x33, 0x33)

_INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|__[^_]+__)")
_NUMBERED_RE = re.compile(r"^\d+[.)]\s+")
_SEPARATOR_RE = re.compile(r":?-{3,}:?")


def _is_table_separator(cells: list[str]) -> bool:
    return bool(cells) and all(_SEPARATOR_RE.fullmatch(c.strip()) for c in cells)


def _add_inline_runs(paragraph, text: str) -> None:
    """Render basic inline Markdown (bold/italic/code) as Word runs."""
    for part in _INLINE_RE.split(text):
        if not part:
            continue
        if (part.startswith("**") and part.endswith("**") and len(part) > 4) or (
            part.startswith("__") and part.endswith("__") and len(part) > 4
        ):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def _add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=col_count)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for c_idx in range(col_count):
            text = row[c_idx] if c_idx < len(row) else ""
            cell = cells[c_idx]
            cell.text = ""
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(text)
            run.font.size = Pt(10)
            if r_idx == 0:
                run.bold = True
                run.font.color.rgb = BLUE
            else:
                run.font.color.rgb = CHARCOAL


def markdown_to_word(markdown_text: str, title: str | None = None) -> Document:
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = CHARCOAL
    normal.paragraph_format.space_after = Pt(6)

    if title:
        heading = doc.add_heading(title, level=0)
        for run in heading.runs:
            run.font.color.rgb = NAVY

    lines = (markdown_text or "").split("\n")
    i = 0
    while i < len(lines):
        stripped = lines[i].strip()

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=2)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=1)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=0)
        elif stripped.startswith(("- ", "* ", "+ ")):
            _add_inline_runs(doc.add_paragraph(style="List Bullet"), stripped[2:])
        elif _NUMBERED_RE.match(stripped):
            _add_inline_runs(
                doc.add_paragraph(style="List Number"),
                _NUMBERED_RE.sub("", stripped),
            )
        elif stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            parsed_rows = []
            for line in table_lines:
                cells = [c.strip() for c in line.strip("|").split("|")]
                parsed_rows.append(cells)
            parsed_rows = [r for r in parsed_rows if not _is_table_separator(r)]
            _add_table(doc, parsed_rows)
            continue
        elif stripped.startswith(">"):
            paragraph = doc.add_paragraph()
            _add_inline_runs(paragraph, stripped.lstrip(">").strip())
            paragraph.paragraph_format.left_indent = Inches(0.3)
            for run in paragraph.runs:
                run.italic = True
        elif stripped in {"---", "***", "___"} or set(stripped) <= {"-", "*", "_"}:
            pass
        elif stripped:
            _add_inline_runs(doc.add_paragraph(), stripped)

        i += 1

    # Style headings navy
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith("Heading"):
            for run in paragraph.runs:
                run.font.color.rgb = NAVY

    return doc


def markdown_to_docx_bytes(markdown_text: str, title: str | None = None) -> bytes:
    document = markdown_to_word(markdown_text, title)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
